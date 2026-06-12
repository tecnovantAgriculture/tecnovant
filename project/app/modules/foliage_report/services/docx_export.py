"""DOCX export service for foliage reports — 9-section executive layout.

Mirrors the client-side PDF export (``report.js::exportDetailedPdf``)
section-for-section:

  1. Portada               — clean cover, no technical data
  2. Resumen Ejecutivo     — KPI dashboard cards
  3. Hallazgos Principales — strengths, risks, priority
  4. Ley del Mínimo        — prominent Liebig block + table
  5. Análisis Foliar       — zebra-striped nutrient table
  6. Macronutrientes       — chart + detailed values
  7. Micronutrientes       — chart + detailed values
  8. Recomendaciones       — Balance de Minerales — Producto Nano
  9. Histórico y Tendencias — chart + trend cards with ▲▼

Charts are rendered server-side with matplotlib (Agg backend).
Multi-tenant: the calling endpoint MUST validate access before calling
:func:`build_report_docx_bytes`.
"""

# Standard library
import io
import os
import unicodedata
from datetime import datetime
from pathlib import Path

# Third party
import matplotlib
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from flask import current_app

# Local application
from app.extensions import db
from app.modules.agrovista.helpers import compute_mineral_balance
from app.modules.foliage.models import Nutrient, ProductPrice, Recommendation
from app.modules.foliage_report.controller import ReportView
from app.modules.foliage_report.helpers import (
    compute_nano_dose_rows,
    contribuciones_de_producto,
    determinar_coeficientes_variacion,
    precios_de_producto,
)
from app.modules.foliage_report.models import RecommendationDose
from app.modules.media.models import StorageLocation

# Brand palette mirrors the PDF jsPDF colors.
BRAND_GREEN = RGBColor(0x10, 0xB9, 0x81)
BRAND_DARK_GREEN = RGBColor(0x06, 0x5F, 0x46)
SOFT_GRAY = RGBColor(0x6B, 0x72, 0x80)
DARK_TEXT = RGBColor(0x11, 0x18, 0x27)
LIGHT_TEXT = RGBColor(0x9C, 0xA3, 0xAF)
TABLE_HEADER_FILL = "065F46"  # emerald-900 for headers
TABLE_ZEBRA_FILL = "F9FAFB"
SECTION_FONT = "Calibri"

# Traffic-light colors
RED_COLOR = RGBColor(0xDC, 0x26, 0x26)
AMBER_COLOR = RGBColor(0xF5, 0x9E, 0x0B)
GREEN_COLOR = RGBColor(0x16, 0xA3, 0x4A)
BLUE_COLOR = RGBColor(0x3B, 0x82, 0xF6)

CHART_COLORS = [
    "#10B981",
    "#3B82F6",
    "#F59E0B",
    "#EF4444",
    "#8B5CF6",
    "#06B6D4",
    "#F97316",
    "#84CC16",
    "#EC4899",
    "#6366F1",
]


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def build_report_docx_bytes(report_id: int) -> io.BytesIO:
    """Build the Word document matching the 9-section PDF layout.

    Args:
        report_id: primary key of ``Recommendation``.

    Returns:
        ``io.BytesIO`` containing a valid ``.docx`` file.

    Raises:
        werkzeug.exceptions.NotFound: when the recommendation does not exist.
    """
    payload = _load_report_payload(report_id)
    cv_data = _load_cv_data(payload)
    mineral_balance = _load_mineral_balance(payload)
    lot_image = _load_lot_image_path(report_id)

    doc = Document()
    _configure_document(doc)

    # 1. Portada
    _render_cover(doc, payload, lot_image)
    doc.add_page_break()

    # 2. Resumen Ejecutivo
    _render_executive_summary(doc, payload, mineral_balance, lot_image)
    doc.add_page_break()

    # 3. Hallazgos Principales
    _render_key_findings(doc, payload)
    doc.add_page_break()

    # 4. Ley del Mínimo de Liebig
    _render_liebig(doc, payload, cv_data)

    # 5. Análisis Foliar Detallado
    doc.add_page_break()
    _render_foliar_detail(doc, payload)

    # 6. Macronutrientes
    _render_nutrient_page(
        doc,
        title="Macronutrientes",
        nutrients=_separate_by_type(payload, "Macronutrient"),
        chart_payload=_build_macro_chart(payload),
        accent=BRAND_GREEN,
    )

    # 7. Micronutrientes
    _render_nutrient_page(
        doc,
        title="Micronutrientes",
        nutrients=_separate_by_type(payload, "Micronutrient"),
        chart_payload=_build_micro_chart(payload),
        accent=RGBColor(0x3B, 0x82, 0xF6),
    )

    # 8. Recomendaciones (Balance de Minerales)
    _render_recommendations(doc, payload, mineral_balance)

    # 9. Histórico y Tendencias
    historical = payload.get("historicalData") or []
    if len(historical) > 1:
        _render_historical(
            doc,
            historical=historical,
            trends=payload.get("trends") or {},
            chart_payload=_build_historical_chart(historical),
        )

    _apply_footer(doc, payload)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Data loading
# ---------------------------------------------------------------------------


def _load_report_payload(report_id: int) -> dict:
    """Reuse the same JSON the web dashboard renders."""
    response = ReportView().get(report_id)
    payload = response.get_json()
    payload["trends"] = _compute_trends(payload.get("historicalData") or [])
    return payload


def _load_recommendation_doses(report_id: int) -> list:
    """Load recommendation doses from DB."""
    rows = (
        RecommendationDose.query.filter_by(recommendation_id=report_id)
        .order_by(RecommendationDose.id.asc())
        .all()
    )
    return [d.to_dict() for d in rows]


def _load_cv_data(payload: dict) -> dict:
    """Replicate CV lookup from vista_reporte."""
    lot = payload.get("lot") or {}
    lot_id = lot.get("id")
    if not lot_id:
        return {}
    try:
        raw = determinar_coeficientes_variacion(lot_id)
        return {k: float(v) for k, v in raw.items()}
    except Exception:
        return {}


def _compute_trends(historical_data: list) -> dict:
    """Recompute trends server-side (mirrors web_routes.calculate_trends)."""
    if not historical_data or len(historical_data) < 2:
        return {}

    def _to_float(value):
        try:
            return float(value)
        except (TypeError, ValueError):
            return None

    def _diff_months(start_iso: str, end_iso: str) -> float:
        try:
            start = datetime.fromisoformat(start_iso)
            end = datetime.fromisoformat(end_iso)
        except (TypeError, ValueError):
            return 1.0
        return max((end - start).days / 30.0, 1.0)

    all_keys = set()
    for entry in historical_data:
        for k in entry:
            if k != "fecha":
                all_keys.add(k)

    trends = {}
    for nutrient in all_keys:
        initial = final = None
        for entry in historical_data:
            val = _to_float(entry.get(nutrient))
            if val is None:
                continue
            if initial is None:
                initial = val
            final = val
        if initial is None or final is None or initial == 0:
            continue
        pct_change = ((final - initial) / initial) * 100.0
        months = _diff_months(
            str(historical_data[0].get("fecha", "")),
            str(historical_data[-1].get("fecha", "")),
        )
        trends[nutrient] = {
            "initial_value": initial,
            "final_value": final,
            "percentage_change": pct_change,
            "monthly_change": pct_change / months,
        }
    return trends


def _load_lot_image_path(report_id: int) -> str | None:
    """Resolve the lot orthophoto file path for embedding in the DOCX.

    Mirrors ``resolve_lot_snapshot_url`` but returns a local filesystem
    path (or None) instead of a URL, so ``doc.add_picture()`` can use it.
    Only local storage is supported — S3 images are skipped.
    """
    rec = Recommendation.query.get(report_id)
    if not rec or not rec.base_analysis:
        return None
    variant = getattr(rec.base_analysis, "lot_snapshot_variant", None)
    if variant is None or variant.storage != StorageLocation.LOCAL.value:
        return None
    try:
        config_dir = current_app.config.get("MEDIA_STORAGE_DIR")
        if not config_dir:
            project_root = os.path.abspath(
                os.path.join(current_app.root_path, os.pardir)
            )
            storage_root = os.path.join(project_root, "storage", "media")
        else:
            storage_root = config_dir
        path = Path(storage_root) / variant.storage_key
    except Exception:
        return None
    return str(path) if path.exists() else None


def _load_mineral_balance(payload: dict) -> dict:
    """Compute mineral balance mirroring web_routes logic.

    Uses foliar data + productive objective from the payload, queries
    Nutrient model for ordering and type classification.
    """
    foliar_data = (payload.get("analysisData") or {}).get("foliar") or {}
    productive_obj = payload.get("productiveObjective") or {}

    name_to_symbol = {
        n.name.lower().replace(" ", ""): n.symbol
        for n in Nutrient.query.order_by(Nutrient.id.asc()).all()
    }
    nutrients = Nutrient.query.order_by(Nutrient.id.asc()).all()
    symbol_position = {n.symbol: i for i, n in enumerate(nutrients)}

    order = []
    targets = {}
    actuals = {}
    for key, entry in foliar_data.items():
        if key == "id" or not isinstance(entry, dict):
            continue
        symbol = name_to_symbol.get(key)
        if not symbol:
            continue
        order.append(symbol)
        if entry.get("ideal") is not None:
            targets[symbol] = float(entry["ideal"])
        if entry.get("valor") is not None:
            actuals[symbol] = float(entry["valor"])

    order.sort(key=lambda s: symbol_position.get(s, 9999))

    aforo = productive_obj.get("target", {}).get("yield") or productive_obj.get(
        "current", {}
    ).get("yield")

    if not order or not aforo:
        return {}

    try:
        balance = compute_mineral_balance(order, targets, actuals, aforo, nutrients)
    except Exception:
        return {}

    # Cost nano doses (best-effort)
    if balance.get("entries"):
        try:
            now = datetime.now()
            price_units = {
                pp.product.name: pp.price_unit
                for pp in ProductPrice.query.filter(
                    ProductPrice.start_date <= now,
                    ProductPrice.end_date >= now,
                ).all()
            }
            balance["nano_doses"] = compute_nano_dose_rows(
                balance,
                {n.symbol: n.name for n in nutrients},
                contribuciones_de_producto(),
                precios_de_producto(),
                price_units,
            )
        except Exception:
            balance["nano_doses"] = {}
    return balance


def _analyze_findings(foliar_data: dict, mla: dict) -> dict:
    """Auto-generate findings from foliar data (mirrors report.js _analyzeFindings)."""
    strengths = []
    risks = []
    limiting = (mla or {}).get("nutriente_limitante")
    limiting_pct = None

    for nut, d in foliar_data.items():
        if nut == "id" or not isinstance(d, dict):
            continue
        ideal = d.get("ideal")
        actual = d.get("valor")
        if not ideal or not actual:
            continue
        pct = (actual / ideal) * 100.0

        if 95 <= pct <= 110:
            strengths.append({"name": nut.capitalize(), "pct": round(pct)})
        elif pct < 80:
            risks.append(
                {"name": nut.capitalize(), "pct": round(pct), "severity": "Deficiencia"}
            )
        elif pct > 140:
            risks.append(
                {"name": nut.capitalize(), "pct": round(pct), "severity": "Exceso"}
            )
        elif pct < 95:
            risks.append(
                {"name": nut.capitalize(), "pct": round(pct), "severity": "Leve"}
            )

        if limiting and nut.lower() == str(limiting).lower():
            limiting_pct = round(pct)

    deficient = sum(1 for r in risks if r["severity"] == "Deficiencia")
    if deficient >= 3 or (limiting_pct is not None and limiting_pct < 70):
        priority, pcolor = "ALTA", RED_COLOR
    elif deficient >= 1 or len(risks) >= 3:
        priority, pcolor = "MEDIA", AMBER_COLOR
    else:
        priority, pcolor = "BAJA", GREEN_COLOR

    return {
        "strengths": strengths,
        "risks": risks,
        "limiting": limiting,
        "limiting_pct": limiting_pct,
        "priority": priority,
        "priority_color": pcolor,
    }


# ---------------------------------------------------------------------------
# Document-level setup
# ---------------------------------------------------------------------------


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)


def _set_cell_shading(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _add_section_heading(
    doc: Document, text: str, color: RGBColor = BRAND_GREEN
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = color
    run.font.name = SECTION_FONT


def _add_body_paragraph(
    doc: Document,
    text: str,
    *,
    bold: bool = False,
    color: RGBColor = DARK_TEXT,
    size: int = 9,
    align=None,
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = SECTION_FONT


def _add_horizontal_line(paragraph, color: RGBColor = BRAND_GREEN) -> None:
    """Add a thin horizontal line (bottom border) to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _traffic_light_color(pct: float) -> RGBColor:
    """Traffic-light color matching PDF jsPDF getTrafficLightColor."""
    if pct < 80:
        return RED_COLOR
    if pct < 95:
        return AMBER_COLOR
    if pct <= 110:
        return GREEN_COLOR
    return BLUE_COLOR


# ---------------------------------------------------------------------------
# 1. Portada
# ---------------------------------------------------------------------------


def _render_cover(doc: Document, payload: dict, lot_image: str | None = None) -> None:
    """Clean executive cover with lot image watermark — no tables, no technical data."""
    com = (payload.get("analysisData") or {}).get("common") or {}
    crop = (payload.get("crop") or {}).get("name") or "--"

    # Decorative lines (top + bottom)
    line_p = doc.add_paragraph()
    line_p.paragraph_format.space_before = Pt(30)
    line_p.paragraph_format.space_after = Pt(0)
    _add_horizontal_line(line_p, BRAND_GREEN)

    # Title block
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(30)
    run = title_p.add_run("Informe de Análisis")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = BRAND_DARK_GREEN
    run.font.name = SECTION_FONT

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = sub_p.add_run("y Recomendaciones")
    sub.font.size = Pt(26)
    sub.font.bold = True
    sub.font.color.rgb = BRAND_DARK_GREEN
    sub.font.name = SECTION_FONT

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline.paragraph_format.space_before = Pt(10)
    tr = tagline.add_run("Ley del Mínimo de Liebig · Nutrición Foliar de Precisión")
    tr.font.size = Pt(10)
    tr.font.italic = True
    tr.font.color.rgb = BRAND_GREEN
    tr.font.name = SECTION_FONT

    # Info block
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(24)

    items = [
        ("Finca", com.get("finca", "--")),
        ("Lote", com.get("lote", "--")),
        ("Cultivo", crop),
        ("Fecha de análisis", com.get("fechaAnalisis", "--")),
        ("Autor", payload.get("author") or "Sistema"),
    ]
    for label, value in items:
        if not value or value == "--":
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        r1 = p.add_run(f"{label}: ")
        r1.font.size = Pt(9)
        r1.font.color.rgb = BRAND_GREEN
        r1.font.name = SECTION_FONT
        r2 = p.add_run(str(value))
        r2.font.size = Pt(9)
        r2.font.bold = True
        r2.font.color.rgb = DARK_TEXT
        r2.font.name = SECTION_FONT

    # Lot image (centered, between info and footer)
    if lot_image and os.path.exists(lot_image):
        try:
            img_spacer = doc.add_paragraph()
            img_spacer.paragraph_format.space_before = Pt(20)
            doc.add_picture(lot_image, width=Inches(4.5))
            last = doc.paragraphs[-1]
            last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

    # Bottom line
    line_bottom = doc.add_paragraph()
    line_bottom.paragraph_format.space_before = Pt(60)
    line_bottom.paragraph_format.space_after = Pt(0)
    _add_horizontal_line(line_bottom, BRAND_GREEN)

    # Date and branding
    fecha = datetime.now().strftime("%d de %B de %Y")
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot.paragraph_format.space_before = Pt(12)
    fr = foot.add_run(fecha)
    fr.font.size = Pt(9)
    fr.font.italic = True
    fr.font.color.rgb = BRAND_GREEN
    fr.font.name = SECTION_FONT

    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = brand.add_run("TecnoAgro · Nutrición Foliar de Precisión")
    br.font.size = Pt(9)
    br.font.color.rgb = SOFT_GRAY
    br.font.name = SECTION_FONT


# ---------------------------------------------------------------------------
# 2. Resumen Ejecutivo
# ---------------------------------------------------------------------------


def _render_executive_summary(
    doc: Document, payload: dict, mineral_balance: dict, lot_image: str | None = None
) -> None:
    """KPI dashboard cards + orthophoto — executive snapshot."""
    _add_section_heading(doc, "Resumen Ejecutivo")

    obj = payload.get("productiveObjective") or {}
    mla = payload.get("minimumLawAnalyses") or {}
    lot = payload.get("lot") or {}
    com = (payload.get("analysisData") or {}).get("common") or {}
    findings = _analyze_findings(
        (payload.get("analysisData") or {}).get("foliar") or {},
        mla,
    )

    # KPI table: 3 columns × 2 rows
    kpi_table = doc.add_table(rows=2, cols=3)
    kpi_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kpi_table.autofit = True

    def _kpi(cell, title, value, subtitle, accent):
        cell.paragraphs[0].clear()
        # Title
        tp = cell.paragraphs[0]
        tr = tp.add_run(title)
        tr.font.size = Pt(7)
        tr.font.bold = True
        tr.font.color.rgb = SOFT_GRAY
        tr.font.name = SECTION_FONT
        # Value
        vp = cell.add_paragraph()
        vr = vp.add_run(value)
        vr.font.size = Pt(14)
        vr.font.bold = True
        vr.font.color.rgb = DARK_TEXT
        vr.font.name = SECTION_FONT
        # Subtitle
        if subtitle:
            sp = cell.add_paragraph()
            sr = sp.add_run(subtitle)
            sr.font.size = Pt(7)
            sr.font.color.rgb = SOFT_GRAY
            sr.font.name = SECTION_FONT
        # Accent line via top border
        _set_cell_shading(cell, "FFFFFF")
        # Set top border via XML
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        top_border = OxmlElement("w:top")
        top_border.set(qn("w:val"), "single")
        top_border.set(qn("w:sz"), "12")
        top_border.set(qn("w:color"), f"{accent[0]:02X}{accent[1]:02X}{accent[2]:02X}")
        borders.append(top_border)
        tc_pr.append(borders)

    # Row 1
    r1 = kpi_table.rows[0].cells
    yield_val = "—"
    yield_sub = ""
    if obj.get("current", {}).get("yield"):
        yield_val = f"{obj['current']['yield']:.1f} t/ha"
    if obj.get("target", {}).get("yield"):
        yield_sub = f"Meta: {obj['target']['yield']:.1f} t/ha"
    _kpi(r1[0], "AFORO", yield_val, yield_sub, (0x10, 0xB9, 0x81))

    prot_val = "—"
    prot_sub = ""
    if obj.get("current", {}).get("protein"):
        prot_val = f"{obj['current']['protein']:.1f}%"
    if obj.get("target", {}).get("protein"):
        prot_sub = f"Meta: {obj['target']['protein']:.1f}%"
    _kpi(r1[1], "PROTEÍNA", prot_val, prot_sub, (0x3B, 0x82, 0xF6))

    lim_name = findings["limiting"] or "Ninguno"
    lim_sub = (
        f"{findings['limiting_pct']}% del ideal"
        if findings["limiting_pct"]
        else "Sin limitante crítico"
    )
    lim_accent = (0xDC, 0x26, 0x26) if findings["limiting"] else (0x16, 0xA3, 0x4A)
    _kpi(r1[2], "LIMITANTE", lim_name, lim_sub, lim_accent)

    # Row 2
    r2 = kpi_table.rows[1].cells
    area_val = f"{lot.get('area', 0):.2f} ha" if lot.get("area") else "—"
    _kpi(
        r2[0],
        "ÁREA DEL LOTE",
        area_val,
        f"Lote {com.get('lote', '—')}",
        (0x6B, 0x72, 0x80),
    )

    estado = "Requiere atención" if findings["risks"] else "Óptimo"
    estado_sub = (
        f"{len(findings['risks'])} fuera de rango"
        if findings["risks"]
        else "Todos en rango"
    )
    estado_accent = (0xF5, 0x9E, 0x0B) if findings["risks"] else (0x16, 0xA3, 0x4A)
    _kpi(r2[1], "ESTADO", estado, estado_sub, estado_accent)

    gaps = obj.get("gaps") or {}
    gap_y = (
        f"Aforo: {gaps['yield_pct']:+.1f}%"
        if gaps.get("yield_pct") is not None
        else "Aforo: —"
    )
    gap_p = (
        f"Proteína: {gaps['protein_pct']:+.1f}%"
        if gaps.get("protein_pct") is not None
        else "Proteína: —"
    )
    _kpi(r2[2], "BRECHA PRODUCTIVA", "", f"{gap_y}  |  {gap_p}", (0xEA, 0x58, 0x0C))

    # Mineral Balance mini-summary
    if mineral_balance.get("total_kg_ha"):
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(
            f"Requerimiento total: {mineral_balance['total_kg_ha']:.2f} kg/ha "
            f"(Producto Nano)"
        )
        r.font.size = Pt(9)
        r.font.bold = True
        r.font.color.rgb = BRAND_GREEN
        r.font.name = SECTION_FONT

    # Orthophoto / NDVI
    if lot_image and os.path.exists(lot_image):
        try:
            img_spacer = doc.add_paragraph()
            img_spacer.paragraph_format.space_before = Pt(8)
            _add_body_paragraph(
                doc, "ORTOFOTO / NDVI", bold=True, color=SOFT_GRAY, size=7
            )
            doc.add_picture(lot_image, width=Inches(5.5))
            last = doc.paragraphs[-1]
            last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass


# ---------------------------------------------------------------------------
# 3. Hallazgos Principales
# ---------------------------------------------------------------------------


def _render_key_findings(doc: Document, payload: dict) -> None:
    """Auto-generated findings: strengths, risks, priority."""
    _add_section_heading(doc, "Hallazgos Principales")

    foliar = (payload.get("analysisData") or {}).get("foliar") or {}
    mla = payload.get("minimumLawAnalyses") or {}
    findings = _analyze_findings(foliar, mla)

    # Priority badge
    badge = doc.add_paragraph()
    badge.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    br = badge.add_run(f"PRIORIDAD: {findings['priority']}")
    br.font.size = Pt(9)
    br.font.bold = True
    br.font.color.rgb = findings["priority_color"]
    br.font.name = SECTION_FONT

    # Strengths
    if findings["strengths"]:
        doc.add_paragraph()
        _add_body_paragraph(doc, "✓ FORTALEZAS", bold=True, color=GREEN_COLOR, size=10)
        items = [f"{s['name']} al {s['pct']}% del ideal" for s in findings["strengths"]]
        _add_body_paragraph(doc, " · ".join(items), size=8)

    # Risks
    if findings["risks"]:
        doc.add_paragraph()
        _add_body_paragraph(doc, "⚠ RIESGOS", bold=True, color=RED_COLOR, size=10)
        for r in findings["risks"]:
            sev = f" ({r['severity']})" if r["severity"] != "Leve" else ""
            _add_body_paragraph(
                doc, f"• {r['name']}{sev} — {r['pct']}% del ideal", size=8
            )

    # Impact
    doc.add_paragraph()
    if findings["limiting"]:
        impact = (
            f"Corregir {findings['limiting']} puede destrabar el potencial "
            "productivo del cultivo. Los nutrientes por debajo del 80% del "
            "ideal están limitando activamente el rendimiento y la calidad."
        )
    else:
        impact = (
            "Los nutrientes se encuentran dentro de rangos aceptables. "
            "Mantener el programa de fertilización actual y monitorear "
            "periódicamente."
        )
    _add_body_paragraph(doc, f"Impacto esperado: {impact}", size=8, color=DARK_TEXT)


# ---------------------------------------------------------------------------
# 4. Ley del Mínimo de Liebig
# ---------------------------------------------------------------------------


def _render_liebig(doc: Document, payload: dict, cv_data: dict) -> None:
    """Prominent Liebig visual block + nutrient summary table."""
    doc.add_page_break()
    _add_section_heading(doc, "Ley del Mínimo de Liebig")

    mla = payload.get("minimumLawAnalyses") or {}
    foliar = (payload.get("analysisData") or {}).get("foliar") or {}
    cv_lookup = {_normalize_key(k): v for k, v in (cv_data or {}).items()}
    findings = _analyze_findings(foliar, mla)

    # Big visual block for limiting nutrient
    lim_name = findings["limiting"] or "Ninguno"
    block = doc.add_paragraph()
    block.alignment = WD_ALIGN_PARAGRAPH.CENTER
    block.paragraph_format.space_before = Pt(8)
    block.paragraph_format.space_after = Pt(8)
    r1 = block.add_run("NUTRIENTE LIMITANTE\n")
    r1.font.size = Pt(9)
    r1.font.color.rgb = BRAND_GREEN
    r1.font.name = SECTION_FONT
    r2 = block.add_run(lim_name)
    r2.font.size = Pt(20)
    r2.font.bold = True
    r2.font.color.rgb = BRAND_DARK_GREEN
    r2.font.name = SECTION_FONT
    if findings["limiting_pct"]:
        r3 = block.add_run(f"\n{findings['limiting_pct']}% del ideal")
        r3.font.size = Pt(9)
        r3.font.color.rgb = BRAND_GREEN
        r3.font.name = SECTION_FONT

    _add_body_paragraph(
        doc,
        "El crecimiento del cultivo está limitado por el nutriente más escaso, "
        "como la duela más corta de un barril determina el nivel de agua.",
        color=SOFT_GRAY,
        size=8,
    )

    # Nutrient summary table
    keys = [k for k in foliar if k != "id" and isinstance(foliar[k], dict)]
    if not keys:
        _add_body_paragraph(doc, "No hay datos foliares disponibles.")
        return

    headers = ["Nutriente", "% Ideal", "I", "R", "Diferencia"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = True
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_shading(hdr[i], TABLE_HEADER_FILL)
        para = hdr[i].paragraphs[0]
        run = para.add_run(h)
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = SECTION_FONT

    for idx, key in enumerate(keys):
        row = table.add_row().cells
        if idx % 2 == 0:
            for c in row:
                _set_cell_shading(c, TABLE_ZEBRA_FILL)
        d = foliar[key]
        actual = d.get("valor")
        target = d.get("ideal")
        if actual is None or target in (None, 0):
            p_val = i_val = r_val = diff_val = None
        else:
            p_val = (actual / target) * 100.0
            norm = _normalize_key(key)
            cv = cv_lookup.get(norm)
            i_val = None if cv is None else abs(p_val - 100.0) * cv / 100.0
            if i_val is None:
                r_val = None
            else:
                r_val = p_val + i_val if p_val < 100 else p_val - i_val
                r_val = max(min(r_val, 108.0), 88.0)
            diff_val = 100.0 - p_val

        display = _short_name(key)
        _write_cell(row[0], display, color=DARK_TEXT)
        _write_cell(
            row[1],
            p_val,
            color=_traffic_light_color(p_val) if p_val else DARK_TEXT,
            fmt="{:.0f}%",
        )
        _write_cell(row[2], i_val, color=DARK_TEXT, fmt="{:.1f}")
        _write_cell(row[3], r_val, color=DARK_TEXT, fmt="{:.1f}")
        diff_color = (
            RED_COLOR
            if (diff_val is not None and diff_val > 0)
            else (AMBER_COLOR if (diff_val is not None and diff_val < 0) else DARK_TEXT)
        )
        _write_cell(row[4], diff_val, color=diff_color, fmt="{:+.0f}%")


# ---------------------------------------------------------------------------
# 5. Análisis Foliar Detallado
# ---------------------------------------------------------------------------


def _render_foliar_detail(doc: Document, payload: dict) -> None:
    """Zebra-striped foliar nutrient table with corporate header."""
    _add_section_heading(doc, "Análisis Foliar Detallado")

    foliar = (payload.get("analysisData") or {}).get("foliar") or {}
    keys = [k for k in foliar if k != "id" and isinstance(foliar[k], dict)]
    if not keys:
        _add_body_paragraph(doc, "No hay datos foliares detallados.")
        return

    headers = ["Nutriente", "Símbolo", "Actual", "Unidad", "Ideal", "Tipo", "% Ideal"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_shading(hdr[i], TABLE_HEADER_FILL)
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.font.size = Pt(7)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = SECTION_FONT

    for idx, key in enumerate(keys):
        row = table.add_row().cells
        if idx % 2 == 0:
            for c in row:
                _set_cell_shading(c, TABLE_ZEBRA_FILL)
        d = foliar[key]
        actual = d.get("valor")
        ideal = d.get("ideal")
        unit = d.get("unidad") or ""
        tipo = (d.get("tipo") or "")[:4]
        symbol = _short_name(key)
        try:
            pct = (actual / ideal) * 100.0 if (actual is not None and ideal) else None
        except (TypeError, ZeroDivisionError):
            pct = None

        _write_cell(row[0], key.capitalize(), color=DARK_TEXT)
        _write_cell(row[1], symbol, color=DARK_TEXT)
        _write_cell(row[2], "" if actual is None else f"{actual}", color=DARK_TEXT)
        _write_cell(row[3], unit, color=DARK_TEXT)
        _write_cell(row[4], "" if ideal is None else f"{ideal}", color=DARK_TEXT)
        _write_cell(row[5], tipo, color=DARK_TEXT)
        if pct is None:
            _write_cell(row[6], "—", color=DARK_TEXT)
        else:
            color = _traffic_light_color(pct)
            _write_cell(row[6], f"{pct:.0f}%", color=color)

    # Soil data (inline, same as PDF foliar page)
    soil = (payload.get("analysisData") or {}).get("soil") or {}
    soil_keys = [k for k in soil if k != "id" and soil[k] not in (None, "")]
    if soil_keys:
        doc.add_paragraph()
        _add_body_paragraph(
            doc, "Análisis de Suelo", bold=True, color=BRAND_GREEN, size=10
        )
        for key in soil_keys:
            _add_body_paragraph(doc, f"{key.capitalize()}: {soil[key]}", size=8)


# ---------------------------------------------------------------------------
# 6-7. Macro / Micro nutrient pages
# ---------------------------------------------------------------------------


def _render_nutrient_page(
    doc: Document, *, title: str, nutrients: dict, chart_payload, accent: RGBColor
) -> None:
    doc.add_page_break()
    _add_section_heading(doc, title, color=accent)

    if chart_payload is not None:
        try:
            doc.add_picture(chart_payload, width=Inches(6.0))
            last = doc.paragraphs[-1]
            last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

    if not nutrients:
        _add_body_paragraph(doc, "Sin datos para esta categoría.")
        return

    sub = doc.add_paragraph()
    sr = sub.add_run("Valores detallados")
    sr.font.size = Pt(11)
    sr.font.bold = True
    sr.font.color.rgb = accent
    sr.font.name = SECTION_FONT
    sub.paragraph_format.space_after = Pt(4)

    for key, d in nutrients.items():
        symbol = _short_name(key)
        actual = d.get("valor")
        ideal = d.get("ideal")
        unit = d.get("unidad") or ""
        try:
            pct = (actual / ideal) * 100.0 if (actual is not None and ideal) else None
        except (TypeError, ZeroDivisionError):
            pct = None
        actual_s = "—" if actual is None else f"{actual:.2f}"
        ideal_s = "—" if ideal is None else f"{ideal:.2f}"
        pct_s = "—" if pct is None else f"{pct:.0f}% del ideal"
        line = f"{symbol}: {actual_s} {unit}    Ideal: {ideal_s} | {pct_s}"
        _add_body_paragraph(doc, line, size=9)


# ---------------------------------------------------------------------------
# 8. Recomendaciones — Balance de Minerales — Producto Nano
# ---------------------------------------------------------------------------


def _render_recommendations(
    doc: Document, payload: dict, mineral_balance: dict
) -> None:
    """Only the Mineral Balance table — no investment summary, no product cards."""
    doc.add_page_break()
    _add_section_heading(doc, "Recomendaciones")

    entries = mineral_balance.get("entries") or []
    total = mineral_balance.get("total_kg_ha")

    if not entries:
        _add_body_paragraph(
            doc,
            "No se pudo calcular el balance de minerales para este reporte.",
            color=SOFT_GRAY,
        )
        return

    # Section title
    _add_body_paragraph(
        doc,
        "Balance de Minerales — Producto Nano",
        bold=True,
        size=10,
        color=BRAND_DARK_GREEN,
    )
    _add_body_paragraph(
        doc,
        "Déficits convertidos a kg/ha según el aforo",
        size=7,
        color=SOFT_GRAY,
    )

    # Build table: Concepto + one column per nutrient + Total
    col_names = ["Concepto"] + [e["name"] for e in entries] + ["Total"]
    rows_data = [
        ("Objetivo (% · ppm)", [e.get("objective_raw") for e in entries], None),
        ("Objetivo (kg/ha)", [e.get("objective_kg") for e in entries], None),
        ("Actual (% · ppm)", [e.get("actual_raw") for e in entries], None),
        ("Actual (kg/ha)", [e.get("actual_kg") for e in entries], None),
        ("Diferencia (kg/ha)", [e.get("difference_kg") for e in entries], None),
        ("Grado fórmula (%)", [e.get("grade_pct") for e in entries], None),
        ("Nano (kg/ha)", [e.get("nano_kg") for e in entries], None),
    ]

    table = doc.add_table(rows=1, cols=len(col_names))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = True

    # Header
    hdr = table.rows[0].cells
    for i, h in enumerate(col_names):
        _set_cell_shading(hdr[i], TABLE_HEADER_FILL)
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.font.size = Pt(7)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = SECTION_FONT

    # Row backgrounds per section
    row_fills = ["EEF2FF", "EEF2FF", "F0F9FF", "F0F9FF", "FEF2F2", "FEFCE8", "ECFDF5"]

    for ri, (label, values, _total) in enumerate(rows_data):
        row = table.add_row().cells
        fill = row_fills[ri % len(row_fills)]
        for c in row:
            _set_cell_shading(c, fill)

        _write_cell(row[0], label, color=DARK_TEXT)
        for ci, v in enumerate(values):
            if v is not None:
                if ri == 4:  # diferencia — red for negative
                    color = RED_COLOR if float(v) < 0 else DARK_TEXT
                elif ri == 6:  # nano — green for positive
                    color = GREEN_COLOR if float(v) > 0 else DARK_TEXT
                else:
                    color = DARK_TEXT
                _write_cell(row[ci + 1], f"{float(v):.2f}", color=color)
            else:
                _write_cell(row[ci + 1], "—", color=DARK_TEXT)

        # Total column
        if ri == 4:
            _write_cell(row[-1], "—", color=DARK_TEXT)
        elif ri == 6:
            _write_cell(row[-1], "—", color=DARK_TEXT)
        elif ri == 5 and total is not None:
            _write_cell(row[-1], f"{total:.2f} kg/ha", color=DARK_TEXT)
        else:
            _write_cell(row[-1], "—", color=DARK_TEXT)

    # Footnote
    doc.add_paragraph()
    _add_body_paragraph(
        doc,
        "Objetivo se expresa en kg/ha con el aforo objetivo y Actual con el "
        "aforo actual (macros: % × aforo × 100; micros: ppm × aforo ÷ 100). "
        "La Diferencia solo refleja déficits. El Grado de fórmula es la "
        "participación porcentual de cada déficit en el total requerido, y "
        "Nano es la cantidad del producto más económico que cubre cada "
        "déficit.",
        size=7,
        color=SOFT_GRAY,
    )
    if mineral_balance.get("aforo_actual_fallback"):
        _add_body_paragraph(
            doc,
            "Nota: no existe aforo actual; la fila Actual (kg/ha) se "
            "convirtió con el aforo objetivo.",
            size=7,
            color=SOFT_GRAY,
        )


# ---------------------------------------------------------------------------
# 9. Histórico y Tendencias
# ---------------------------------------------------------------------------


def _render_historical(
    doc: Document, *, historical: list, trends: dict, chart_payload
) -> None:
    """Historical chart + trend cards with ▲▼ arrows."""
    doc.add_page_break()
    _add_section_heading(doc, "Histórico y Tendencias")

    if chart_payload is not None:
        try:
            doc.add_picture(chart_payload, width=Inches(6.0))
            last = doc.paragraphs[-1]
            last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

    # Trend cards
    if trends:
        doc.add_paragraph()
        _add_body_paragraph(
            doc,
            "Variaciones observadas",
            bold=True,
            color=DARK_TEXT,
            size=10,
        )
        for nutrient, t in trends.items():
            pct = t.get("percentage_change")
            if pct is None:
                continue
            arrow = "▲" if pct >= 0 else "▼"
            color = GREEN_COLOR if pct >= 0 else RED_COLOR
            text = (
                f"{arrow} {nutrient}: {abs(pct):.1f}%  "
                f"({t.get('initial_value', 0):.1f}% → "
                f"{t.get('final_value', 0):.1f}%)"
            )
            _add_body_paragraph(doc, text, size=8, color=color)

    # Historical data table
    nutrient_keys = set()
    for entry in historical:
        for k in entry:
            if k != "fecha":
                nutrient_keys.add(k)
    nutrient_keys = sorted(nutrient_keys)

    if nutrient_keys:
        doc.add_paragraph()
        headers = ["Fecha"] + nutrient_keys
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            _set_cell_shading(hdr[i], TABLE_HEADER_FILL)
            p = hdr[i].paragraphs[0]
            run = p.add_run(h)
            run.font.size = Pt(7)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.name = SECTION_FONT

        for idx, entry in enumerate(historical):
            row = table.add_row().cells
            if idx % 2 == 0:
                for c in row:
                    _set_cell_shading(c, TABLE_ZEBRA_FILL)
            _write_cell(row[0], str(entry.get("fecha", "")), color=DARK_TEXT)
            for j, n in enumerate(nutrient_keys, start=1):
                v = entry.get(n)
                _write_cell(row[j], "" if v in (None, "") else str(v), color=DARK_TEXT)


# ---------------------------------------------------------------------------
# Footer
# ---------------------------------------------------------------------------


def _apply_footer(doc: Document, payload: dict) -> None:
    com = (payload.get("analysisData") or {}).get("common") or {}
    finca = com.get("finca", "")
    lote = com.get("lote", "")
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Informe Agronómico · {finca} · {lote} · TecnoAgro")
    run.font.size = Pt(7)
    run.font.color.rgb = LIGHT_TEXT
    run.font.name = SECTION_FONT


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _write_cell(cell, value, *, color: RGBColor, fmt: str = "{}") -> None:
    para = cell.paragraphs[0]
    if value is None:
        text = "—"
    elif isinstance(value, str):
        text = value
    else:
        try:
            text = fmt.format(value)
        except (TypeError, ValueError):
            text = str(value)
    run = para.add_run(text)
    run.font.size = Pt(7)
    run.font.color.rgb = color
    run.font.name = SECTION_FONT
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _short_name(key: str) -> str:
    return {
        "nitrogeno": "N",
        "fosforo": "P",
        "potasio": "K",
        "calcio": "Ca",
        "magnesio": "Mg",
        "azufre": "S",
        "hierro": "Fe",
        "manganeso": "Mn",
        "zinc": "Zn",
        "cobre": "Cu",
        "boro": "B",
        "molibdeno": "Mo",
        "silicio": "Si",
    }.get(key, key.upper())


def _normalize_key(name: str) -> str:
    return "".join(
        c
        for c in unicodedata.normalize("NFD", name.lower())
        if unicodedata.category(c) != "Mn"
    )


def _separate_by_type(payload: dict, nutrient_type: str) -> dict:
    foliar = (payload.get("analysisData") or {}).get("foliar") or {}
    out = {}
    for k, v in foliar.items():
        if k == "id" or not isinstance(v, dict):
            continue
        if v.get("tipo") == nutrient_type:
            out[k] = v
    return out


# ---------------------------------------------------------------------------
# Matplotlib chart builders
# ---------------------------------------------------------------------------


def _build_progress_bar_chart(nutrients: dict, *, title: str):
    matplotlib.use("Agg")
    items = list(nutrients.items())
    if not items:
        return None
    labels = [_short_name(k) for k, _ in items]
    pcts = []
    for _, d in items:
        actual = d.get("valor")
        ideal = d.get("ideal")
        try:
            pcts.append((actual / ideal) * 100.0 if ideal else 0.0)
        except (TypeError, ZeroDivisionError):
            pcts.append(0.0)

    colors = list(CHART_COLORS)
    if len(pcts) > len(colors):
        for i in range(len(colors), len(pcts)):
            colors.append(_hsl_color(i))

    # Figure width matches DOCX insertion width so no distortion occurs.
    # Height scales with bar count: ~0.4 in per bar + label/title overhead.
    fig_w = 7.0
    fig_h = max(2.8, 0.40 * len(items) + 1.3)

    fig, ax = plt.subplots(figsize=(fig_w, fig_h))
    bar_h = min(0.50, 2.0 / max(len(items), 1))
    bars = ax.barh(
        labels, pcts, color=colors[: len(pcts)], edgecolor="white", height=bar_h
    )
    ax.set_xlim(0, max(150, max(pcts) * 1.15))
    ax.set_xlabel("% del Ideal", fontsize=10, fontweight="bold")
    ax.set_title(title, fontsize=12, fontweight="bold", color="#10B981")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(axis="x", color="#cccccc", linestyle="-", linewidth=0.5)
    ax.set_axisbelow(True)

    for bar, pct in zip(bars, pcts):
        ax.text(
            bar.get_width() + 2,
            bar.get_y() + bar.get_height() / 2,
            f"{pct:.1f}%",
            va="center",
            fontsize=8,
            color="#1F2937",
        )

    # Tight margins without bbox_inches cropping.
    fig.subplots_adjust(left=0.10, right=0.94, top=0.92, bottom=0.10)

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=200)
    plt.close(fig)
    buf.seek(0)
    return buf


def _build_line_chart(historical: list):
    matplotlib.use("Agg")
    if not historical or len(historical) < 2:
        return None
    labels = [str(e.get("fecha", "")) for e in historical]

    nutrient_keys = []
    for entry in historical:
        for k in entry:
            if k != "fecha" and k not in nutrient_keys:
                nutrient_keys.append(k)
    if not nutrient_keys:
        return None

    fig, ax = plt.subplots(figsize=(7.5, 4.0))
    for idx, key in enumerate(nutrient_keys):
        values = []
        for entry in historical:
            try:
                values.append(float(entry.get(key)))
            except (TypeError, ValueError):
                values.append(None)
        cleaned = [v if v is not None else float("nan") for v in values]
        color = CHART_COLORS[idx % len(CHART_COLORS)]
        ax.plot(labels, cleaned, marker="o", linewidth=1.5, label=key, color=color)

    ax.set_title("Evolución Histórica", fontsize=12, fontweight="bold", color="#10B981")
    ax.set_xlabel("Fecha", fontsize=10)
    ax.set_ylabel("Valor", fontsize=10)
    ax.grid(True, color="#cccccc", linestyle="-", linewidth=0.5)
    ax.set_axisbelow(True)
    ax.tick_params(axis="x", labelrotation=30, labelsize=7)
    if nutrient_keys:
        ax.legend(loc="best", fontsize=7, frameon=False)
    plt.tight_layout()
    return _fig_to_buffer(fig)


def _hsl_color(i: int) -> str:
    hue = (i * 137.5) % 360
    return f"hsl({hue}, 70%, 50%)"


def _fig_to_buffer(fig) -> io.BytesIO:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=200, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _build_macro_chart(payload: dict):
    return _build_progress_bar_chart(
        _separate_by_type(payload, "Macronutrient"), title="Macronutrientes"
    )


def _build_micro_chart(payload: dict):
    return _build_progress_bar_chart(
        _separate_by_type(payload, "Micronutrient"), title="Micronutrientes"
    )


def _build_historical_chart(historical: list):
    return _build_line_chart(historical)
